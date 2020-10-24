import re
from os.path import isfile,join
from os import listdir
from django.shortcuts import render
# from django.http import HttpResponse
# from django.views import View
import docx
from docx import Document
from .models import Book
# from docx import Document
# Create your views here.


# Hiện trang

def index_result(request):
    if request.method == 'POST':
        name = request.POST.get('search')

    files = [f for f in listdir('static/docxfile') if isfile(join('static/docxfile',f))]
    # file result
    input_content = []
    res = []
    for file in files:
        document = Document('static//docxfile//' + file)
        fulltext = ''
        for para in document.paragraphs:
            fulltext = fulltext + para.text

        if name in fulltext:
            input_content += [file]
            # print(input_content)         
        else:
            # name = " "
            # print('Xâu tìm kiếm không có trong file '+ file)
            thongbao = "Không có kết quả phù hợp"
            context = {
                'thongbao' : thongbao,
            }
    for i in input_content :  
        noidung = Book.objects.values().filter(noi_dung=i)
        res.append(list(noidung))
        context = {
            'res':res,
            'name':name,
        } 
    # print("================",res)

    return render(request, 'pages/result.html', context)




# Hiện trang với thông tin từ db (thông tin dc lấy là một dictionary)
def display(request):
    # print(virtual)
    book = Book.objects.all() # lấy mọi dòng từ bảng Book trong db
    # context: dictionary bao gồm cặp tên biến & giá trị: {'key': value}
    context = {
        'book': book,
    }
    return render(request, 'pages/base.html',context) #render trang với context

def test(request, id = id):
    t = Book.objects.filter(id = id)
    nd = Book.objects.values('noi_dung').filter(id = id) # lấy tên của sách (cột noi_dung) được chọn theo id
    vals = list(nd.values())[0]['noi_dung']
    print(vals)
    f = docx.Document("static/docxfile/" + vals)
    file_content = []
    for para in f.paragraphs:
        file_content.append(para.text)
    final = '\n'.join(file_content)
    # f.close()
    context = {
        't' : t , 
        'file_content' :file_content ,
        'final': final,
    }
    return render(request, 'pages/test.html', context)
#sửa lỗi class has no 'objects' member
# cmd:  pip install pylint-django
#       Ctrl + shift + P > Preferences: Configure Language Specific Settings > Python
#       sửa file settings.json thành:
#     {
#       "python.linting.pylintArgs": [
#         "--load-plugins=pylint_django"
#       ],

#       "[python]": {

#     }
# }